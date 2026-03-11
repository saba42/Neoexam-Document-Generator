import os
import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import docx
from supabase import create_client, Client

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

def add_run(p, text, bold=False, size=11, color=None):
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run

def add_paragraph(doc, text="", bold=False, size=11, color=None, style=None):
    p = doc.add_paragraph(style=style)
    if text:
        add_run(p, text, bold, size, color)
    return p

def add_watermark(doc):
    from backend.utils.paths import get_data_path
    image_path = get_data_path("watermark.png")
    if not os.path.exists(image_path):
        return

    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
            
        rId, _ = header.part.get_or_add_image(image_path)
        
        xml_str = f'''
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" 
             xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" 
             xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
            <w:rPr>
                <w:noProof/>
            </w:rPr>
            <w:drawing>
                <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
                    <wp:simplePos x="0" y="0"/>
                    <wp:positionH relativeFrom="page">
                        <wp:align>center</wp:align>
                    </wp:positionH>
                    <wp:positionV relativeFrom="page">
                        <wp:align>center</wp:align>
                    </wp:positionV>
                    <wp:extent cx="3657600" cy="1086307"/>
                    <wp:effectExtent l="0" t="0" r="0" b="0"/>
                    <wp:wrapNone/>
                    <wp:docPr id="100" name="Watermark"/>
                    <wp:cNvGraphicFramePr>
                        <a:graphicFrameLocks noChangeAspect="1"/>
                    </wp:cNvGraphicFramePr>
                    <a:graphic>
                        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                            <pic:pic>
                                <pic:nvPicPr>
                                    <pic:cNvPr id="1" name="Watermark" descr="Watermark"/>
                                    <pic:cNvPicPr/>
                                </pic:nvPicPr>
                                <pic:blipFill>
                                    <a:blip r:embed="{rId}">
                                        <a:alphaModFix amount="10000"/>
                                    </a:blip>
                                    <a:stretch>
                                        <a:fillRect/>
                                    </a:stretch>
                                </pic:blipFill>
                                <pic:spPr>
                                    <a:xfrm rot="-2700000">
                                        <a:off x="0" y="0"/>
                                        <a:ext cx="3657600" cy="1086307"/>
                                    </a:xfrm>
                                    <a:prstGeom prst="rect">
                                        <a:avLst/>
                                    </a:prstGeom>
                                </pic:spPr>
                            </pic:pic>
                        </a:graphicData>
                    </a:graphic>
                </wp:anchor>
            </w:drawing>
        </w:r>
        '''
        
        drawing_run = parse_xml(xml_str)
        p = header.paragraphs[0]
        p._p.append(drawing_run)

from storage.supabase_loader import get_source_document_path

def generate_document(parameters, course_code, test_name, output_filename):
    doc = Document()
    set_style(doc)

    import re
    def sanitize_bookmark(text):
        clean = re.sub(r'\W+', '_', text)
        return clean[:40].strip('_')

    def add_internal_hyperlink(p, text, bookmark_name):
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        hyperlink.append(r)
        p._p.append(hyperlink)

    def add_bookmark(p, bookmark_name, b_id):
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), str(b_id))
        start.set(qn('w:name'), bookmark_name)
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), str(b_id))
        p._p.insert(0, start)
        p._p.append(end)

    import importlib.util
    docs_db = {}
    
    try:
        source_path = get_source_document_path()
    except FileNotFoundError as e:
        print(f"❌ {e}")
        source_path = None
        
    if source_path:
        parser_path = os.path.join(os.path.dirname(__file__), 'doc_parser.py')
        if os.path.exists(parser_path):
            spec = importlib.util.spec_from_file_location("doc_parser", parser_path)
            parser = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(parser)
            docs_db = parser.parse_source_document(source_path)
        else:
            print("doc_parser.py not found.")

    # HEADER SECTION
    NAVY_BLUE = RGBColor(0x00, 0x33, 0x66)
    BLACK = RGBColor(0x00, 0x00, 0x00)

    date_str = datetime.datetime.now().strftime("%d %B %Y")
    
    add_paragraph(doc, 'Exam Instructions', bold=True, size=18, color=NAVY_BLUE)
    add_paragraph(doc, f"Course: {course_code}", size=11)
    add_paragraph(doc, f"Test: {test_name}", size=11)
    add_paragraph(doc, f"Date: {date_str}", size=11)
    
    doc.add_paragraph() # spacing
    
    # OVERVIEW SECTION
    add_paragraph(doc, "Examination Configuration Overview", bold=True, size=14)
    add_paragraph(doc, "This document outlines the examination configuration and the enabled system settings for the above-mentioned test. All faculty members and examination staff are requested to carefully review and verify the configured settings prior to the commencement of the examination to ensure accuracy, compliance, and the smooth conduct of the assessment process.", size=11)
    
    doc.add_paragraph() # spacing
    
    # PARAMETER ENABLED LIST
    add_paragraph(doc, "List of Enabled Test Control Parameters", bold=True, size=14)
    add_paragraph(doc, "The examination has been configured with the following control parameters enabled:", size=11)
    
    for param in parameters:
        ui_name = str(param.get("name", "")).strip()
        b_name = sanitize_bookmark(ui_name)
        p_list = doc.add_paragraph(style='List Bullet')
        add_internal_hyperlink(p_list, ui_name, b_name)
        
    doc.add_paragraph() # spacing
    
    # Divider BEFORE parameter documentation
    p_param_div = doc.add_paragraph()
    r_param_div = p_param_div.add_run("════════════════════")
    r_param_div.font.size = Pt(8)
    doc.add_paragraph()

    def normalize_param(text):
        text = str(text).strip().lower()
        text = text.rstrip(":")
        text = text.replace("&", "and")
        text = text.replace("–", "-")
        text = text.replace("—", "-")
        text = text.replace("onetoone", "one to one")
        text = text.replace("enable ", "")
        text = text.replace("disable ", "")
        text = re.sub(r'[^a-z0-9\s]', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def find_param_data(db, checked_name):
        norm = normalize_param(checked_name)
        for cat in db:
            for key in db[cat]:
                norm_key = normalize_param(key)
                if norm == norm_key:
                    return db[cat][key]
                if len(norm) > 4 and norm_key in norm:
                    return db[cat][key]
                if len(norm_key) > 4 and norm in norm_key:
                    return db[cat][key]
                w1 = set(norm.split())
                w2 = set(norm_key.split())
                overlap = len(w1 & w2)
                total = len(w1 | w2)
                if total > 0 and overlap/total >= 0.6:
                    return db[cat][key]
                    
        # Category fallback for top-level toggle parameters
        for cat in db:
            norm_cat = normalize_param(cat)
            if len(norm_cat) > 4 and (norm_cat in norm or norm in norm_cat):
                return {
                    "description_blocks": [
                         {"type": "text", "content": f"The {cat} configuration module has been successfully enabled."}
                    ],
                    "faqs": []
                }
                
        return None

    # PARAMETER DEFINITIONS
    for idx, param in enumerate(parameters, start=1):
        p_name = param.get("name", "").strip()
        b_name = sanitize_bookmark(p_name)
        
        # FEATURE NAME (14pt bold black) WITHOUT idx
        p_head = add_paragraph(doc, p_name, bold=True, size=14, color=BLACK)
        add_bookmark(p_head, b_name, idx)
        doc.add_paragraph()
            
        # Exact + fuzzy match via find_param_data()
        db_info = find_param_data(docs_db, p_name)
            
        db_desc_blocks = db_info.get("description_blocks", []) if db_info else []
        db_faqs = db_info.get("faqs", []) if db_info else []
    
        has_doc = bool(db_desc_blocks or db_faqs)
    
        if not has_doc:
            add_paragraph(doc, "Documentation coming soon", size=11)
        else:
            if db_desc_blocks:
                for block in db_desc_blocks:
                    if block["type"] == "text":
                        txt_content = block["content"].replace("{value}", param.get("value", "")).replace("{interval}", param.get("value", ""))
                        for line in txt_content.split('\n'):
                            line = line.strip()
                            if not line: continue
                        
                            lower_l = line.lower()
                            if lower_l.startswith("definition:"):
                                add_paragraph(doc, "Definition:", bold=True, size=11)
                                rem = line[len("definition:"):].strip()
                                if rem: add_paragraph(doc, rem, size=11)
                            elif lower_l.startswith("how it works:"):
                                add_paragraph(doc, "How It Works:", bold=True, size=11)
                                rem = line[len("how it works:"):].strip()
                                if rem: add_paragraph(doc, rem, size=11)
                            elif line.startswith("-") or line.startswith("•"):
                                clean_line = line.lstrip("-•").strip()
                                if not clean_line:
                                    add_paragraph(doc, line, size=11)
                                else:
                                    add_paragraph(doc, clean_line, size=11, style='List Bullet')
                            else:
                                add_paragraph(doc, line, size=11)
                    elif block["type"] == "image":
                        from io import BytesIO
                        from docx.shared import Inches
                        import traceback
                        try:
                            stream = BytesIO(block["blob"])
                            img_p = doc.add_paragraph()
                            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_p.add_run()
                            run.add_picture(stream, width=Inches(5.0))
                        except Exception as e:
                            pass
                        
                doc.add_paragraph()
                    
            if db_faqs:
                add_paragraph(doc, "FAQ:", bold=True, size=11)
                doc.add_paragraph()
                for fidx, faq in enumerate(db_faqs, start=1):
                    q = faq.get("question", "").strip()
                    q_clean = re.sub(r'^q\d+\s*[:-]?\s*', '', q, flags=re.IGNORECASE).strip()
                
                    add_paragraph(doc, f"Q{fidx}: {q_clean}", bold=True, size=11)
                
                    ans_blocks = faq.get("answer_blocks",[])
                    if ans_blocks:
                        add_paragraph(
                            doc, "Solution:", 
                            bold=True, size=11
                        )
                        for block in ans_blocks:
                            if block["type"] == "text":
                                txt = block["content"]
                                for line in txt.split('\n'):
                                    line = line.strip()
                                    if not line:
                                        continue
                                    if line.lower().startswith(
                                        "solution:"):
                                        rest = line[9:].strip()
                                        if rest:
                                            add_paragraph(
                                                doc, rest,
                                                bold=False, size=11
                                            )
                                    else:
                                        add_paragraph(
                                            doc, line,
                                            bold=False, size=11
                                        )
                            elif block["type"] == "image":
                                from io import BytesIO
                                from docx.shared import Inches
                                try:
                                    stream = BytesIO(block["blob"])
                                    img_p = doc.add_paragraph()
                                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = img_p.add_run()
                                    run.add_picture(
                                        stream, width=Inches(5.0)
                                    )
                                except:
                                    pass
                    else:
                        a = faq.get("answer", "").strip()
                        if a:
                            for line in a.split('\n'):
                                line = line.strip()
                                if line: add_paragraph(doc, line, size=11)
                                
                    if fidx < len(db_faqs):
                        doc.add_paragraph()
    
        # Divider AFTER parameter documentation
        pass

    # FOOTER
    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]

    add_run(footer_p, "Page ", size=9)

    # Page num logic
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "1"
    fldChar2.append(fldChar3)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r = footer_p.add_run()
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar4)

    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not output_filename.endswith(".docx"):
        output_filename += ".docx"

    from backend.utils.paths import get_output_path
    out_filepath = get_output_path(output_filename)

    # Add watermark just before saving
    add_watermark(doc)

    # Save document
    doc.save(out_filepath)
    return out_filepath




